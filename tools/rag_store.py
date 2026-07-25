"""Lightweight updatable RAG store.

- Documents are chunked and saved to data/knowledge_base/docs.json.
- Retrieval uses TF-IDF cosine similarity (scikit-learn), so it works
  identically on local machines and on Streamlit Cloud with no embedding
  API and no GPU. The store can be updated at runtime: add, replace, or
  delete documents and the index rebuilds automatically.
"""
from __future__ import annotations

import json
import re
import time
from pathlib import Path
from typing import Any

KB_DIR = Path(__file__).resolve().parents[1] / "data" / "knowledge_base"
KB_FILE = KB_DIR / "docs.json"

CHUNK_SIZE = 900
CHUNK_OVERLAP = 150


# ---------------------------------------------------------------- storage

def _load_raw() -> list[dict[str, Any]]:
    if not KB_FILE.exists():
        return []
    try:
        return json.loads(KB_FILE.read_text(encoding="utf-8"))
    except Exception:
        return []


def _save_raw(chunks: list[dict[str, Any]]) -> None:
    KB_DIR.mkdir(parents=True, exist_ok=True)
    KB_FILE.write_text(json.dumps(chunks, ensure_ascii=False, indent=1), encoding="utf-8")


# ---------------------------------------------------------------- chunking

def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = re.sub(r"\s+", " ", str(text)).strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        # try to break at a sentence or word boundary
        if end < len(text):
            cut = text.rfind(". ", start + size // 2, end)
            if cut == -1:
                cut = text.rfind(" ", start + size // 2, end)
            if cut != -1:
                end = cut + 1
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return [c for c in chunks if c]


# ---------------------------------------------------------------- public API

def add_document(name: str, text: str, replace: bool = True) -> int:
    """Add or update a document. Returns the number of chunks stored."""
    chunks = _load_raw()
    if replace:
        chunks = [c for c in chunks if c["source"] != name]
    now = time.strftime("%Y-%m-%d %H:%M")
    new = [
        {"source": name, "chunk_id": i, "text": piece, "added_at": now}
        for i, piece in enumerate(chunk_text(text))
    ]
    _save_raw(chunks + new)
    return len(new)


def remove_document(name: str) -> int:
    chunks = _load_raw()
    kept = [c for c in chunks if c["source"] != name]
    _save_raw(kept)
    return len(chunks) - len(kept)


def list_documents() -> list[dict[str, Any]]:
    docs: dict[str, dict[str, Any]] = {}
    for c in _load_raw():
        entry = docs.setdefault(c["source"], {"source": c["source"], "chunks": 0, "added_at": c.get("added_at", "")})
        entry["chunks"] += 1
    return sorted(docs.values(), key=lambda d: d["source"])


def kb_size() -> int:
    return len(_load_raw())


def search(query: str, k: int = 4, min_score: float = 0.05) -> list[dict[str, Any]]:
    """Return the top-k most relevant chunks for the query."""
    chunks = _load_raw()
    if not chunks or not str(query).strip():
        return []
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
    except Exception:
        return _keyword_fallback(query, chunks, k)
    texts = [c["text"] for c in chunks]
    vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), sublinear_tf=True)
    try:
        matrix = vectorizer.fit_transform(texts + [query])
    except ValueError:  # empty vocabulary
        return []
    scores = cosine_similarity(matrix[-1], matrix[:-1]).ravel()
    order = scores.argsort()[::-1][:k]
    results = []
    for idx in order:
        if scores[idx] < min_score:
            continue
        c = chunks[idx]
        results.append({"source": c["source"], "text": c["text"], "score": round(float(scores[idx]), 3)})
    return results


def _keyword_fallback(query: str, chunks: list[dict[str, Any]], k: int) -> list[dict[str, Any]]:
    words = {w for w in re.findall(r"[a-zA-Z]{3,}", query.lower())}
    scored = []
    for c in chunks:
        text_words = set(re.findall(r"[a-zA-Z]{3,}", c["text"].lower()))
        overlap = len(words & text_words)
        if overlap:
            scored.append((overlap, c))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [{"source": c["source"], "text": c["text"], "score": float(s)} for s, c in scored[:k]]


# ---------------------------------------------------------------- file readers

def read_uploaded_file(filename: str, data: bytes) -> str:
    """Extract text from an uploaded file (pdf, docx, txt, md, csv)."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        import io
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    # txt, md, csv, anything else text-like
    return data.decode("utf-8", errors="ignore")
